"""Markdown report to DOCX."""

import io
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown report to DOCX bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for line in markdown_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        # Headers
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_heading(stripped[2:], level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            continue

        # Blockquotes (disclaimer)
        if stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
            continue

        # Horizontal rules
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 50)
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # Numbered lists
        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            doc.add_paragraph(num_match.group(1), style="List Number")
            continue

        # Regular paragraph — strip bold markers for plain text
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
        clean = re.sub(r"\*(.+?)\*", r"\1", clean)
        doc.add_paragraph(clean)

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
